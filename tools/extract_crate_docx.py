"""
Extract images and metadata from crate .docx files.

Outputs:
- data/processed/crate_images/<crate>/<doc_name>/image_*.ext
- data/processed/crate_docs.csv with per-doc metadata
"""

from __future__ import annotations

import csv
import io
import zipfile
from pathlib import Path
from typing import List, Tuple

from docx import Document
from PIL import Image

CRATE_FOLDERS = [
    Path("data/raw/Crate 1 - Completed"),
    Path("data/raw/Crate 2 - Completed"),
    Path("data/raw/Crate 3 - Completed"),
]

OUTPUT_ROOT = Path("data/processed/crate_images")
CSV_PATH = Path("data/processed/crate_docs.csv")


def extract_media(docx_path: Path, output_dir: Path) -> List[str]:
    """Extract images from a .docx into output_dir and return relative paths."""
    output_dir.mkdir(parents=True, exist_ok=True)
    saved = []
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                # Standardize to JPEG filenames
                stem = Path(name).stem
                target = output_dir / f"{stem}.jpg"
                with zf.open(name) as src:
                    data = src.read()
                try:
                    with Image.open(io.BytesIO(data)).convert("RGB") as im:
                        im.save(target, format="JPEG", quality=90)
                    saved.append(str(target))
                except Exception:
                    # Fallback to raw write if Pillow cannot handle format
                    with target.open("wb") as dst:
                        dst.write(data)
                    saved.append(str(target))
    return saved


def parse_table(doc: Document) -> Tuple[str, str, str]:
    """Return (crab_number, carapace_width_mm, color, degree_of_molt) from first table row if present."""
    crab_num = carapace = color = degree = ""
    if not doc.tables:
        return crab_num, carapace, color, degree
    table = doc.tables[0]
    if len(table.rows) < 2:
        return crab_num, carapace, color, degree
    row = table.rows[1].cells
    crab_num = row[0].text.strip()
    carapace = row[1].text.strip()
    color = row[2].text.strip()
    degree = row[3].text.strip() if len(row) > 3 else ""
    return crab_num, carapace, color, degree


def parse_paragraphs(doc: Document) -> List[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def process_docx(docx_path: Path, crate_name: str) -> dict:
    doc = Document(docx_path)
    crab_num, carapace, color, degree = parse_table(doc)
    paragraphs = parse_paragraphs(doc)

    out_dir = OUTPUT_ROOT / crate_name / docx_path.stem
    images = extract_media(docx_path, out_dir)

    return {
        "crate": crate_name,
        "doc_name": docx_path.name,
        "crab_number": crab_num,
        "carapace_width_mm": carapace,
        "color": color,
        "degree_of_molt": degree,
        "paragraphs": "; ".join(paragraphs),
        "image_paths": "|".join(images),
    }


def main():
    records = []
    for crate in CRATE_FOLDERS:
        crate_name = crate.name
        if not crate.exists():
            continue
        for docx_file in crate.glob("*.docx"):
            # Skip non-crab logs
            if docx_file.name.lower().startswith(("observations", "salinity")):
                continue
            rec = process_docx(docx_file, crate_name)
            records.append(rec)

    CSV_PATH.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "crate",
        "doc_name",
        "crab_number",
        "carapace_width_mm",
        "color",
        "degree_of_molt",
        "paragraphs",
        "image_paths",
    ]
    with CSV_PATH.open("w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for rec in records:
            writer.writerow(rec)
    print(f"Wrote {len(records)} records to {CSV_PATH}")


if __name__ == "__main__":
    main()
